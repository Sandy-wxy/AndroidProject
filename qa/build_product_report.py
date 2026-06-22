from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/yuuna/Downloads/Focus_Flow")
SOURCE = ROOT / "选题号-分组-学号-姓名-产品设计书.docx"
OUTPUT = ROOT / "Focus_Flow-产品设计书.docx"

TEAL = RGBColor(0x08, 0x91, 0xB2)
NAVY = RGBColor(0x11, 0x18, 0x27)
GRAY = RGBColor(0x47, 0x55, 0x69)
RED = RGBColor(0xD9, 0x2D, 0x20)
LIGHT_TEAL = "E7F7FA"


def set_run_font(run, size=10.5, bold=False, color=NAVY, font="宋体"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_spacing(paragraph, before=0, after=3, line=1.05):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def clear_body(document):
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "70")
    spacing.set(qn("w:after"), "70")
    p_pr.append(spacing)


def add_heading(document, text):
    p = document.add_paragraph()
    set_spacing(p, before=5, after=3, line=1.0)
    r = p.add_run(text)
    set_run_font(r, size=13.5, bold=True, color=TEAL, font="黑体")
    return p


def add_labeled_paragraph(document, label, text):
    p = document.add_paragraph()
    set_spacing(p, after=2)
    r = p.add_run(label)
    set_run_font(r, bold=True, color=NAVY, font="黑体")
    r = p.add_run(text)
    set_run_font(r, color=GRAY)
    return p


def add_bullet(document, title, detail, innovative=False):
    p = document.add_paragraph(style="List Paragraph")
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "2")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)
    p.paragraph_format.left_indent = Cm(0.65)
    p.paragraph_format.first_line_indent = Cm(-0.28)
    set_spacing(p, after=1.5, line=1.02)
    color = RED if innovative else NAVY
    r = p.add_run(title)
    set_run_font(r, size=10.3, bold=True, color=color, font="黑体")
    r = p.add_run("：" + detail)
    set_run_font(r, size=10.3, color=color if innovative else GRAY)
    return p


doc = Document(SOURCE)
clear_body(doc)

section = doc.sections[0]
section.top_margin = Cm(1.45)
section.bottom_margin = Cm(1.35)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name = "宋体"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(title, after=2, line=1.0)
r = title.add_run("Focus Flow 产品设计书")
set_run_font(r, size=18, bold=True, color=NAVY, font="黑体")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(meta, after=5, line=1.0)
r = meta.add_run("选题号：________    分组：________    学号：________    姓名：________")
set_run_font(r, size=9.5, color=GRAY)

add_heading(doc, "1. 项目简介")
add_labeled_paragraph(doc, "APP 名称：", "Focus Flow（专注流）")
add_labeled_paragraph(
    doc,
    "解决问题：",
    "面向学习任务分散、计划与执行脱节、容易拖延、缺少及时提醒和复盘依据等问题，"
    "把“任务安排—闹钟提醒—番茄专注—数据复盘”整合为一个闭环。",
)
add_labeled_paragraph(
    doc,
    "目标用户：",
    "大学生、备考人群、自主学习者，以及需要按日期规划任务和培养稳定专注习惯的用户。",
)

positioning = doc.add_paragraph()
set_spacing(positioning, before=2, after=4, line=1.05)
shade_paragraph(positioning, LIGHT_TEAL)
r = positioning.add_run("产品定位  ")
set_run_font(r, size=10.3, bold=True, color=TEAL, font="黑体")
r = positioning.add_run(
    "一款离线优先、注重隐私的 Android 学习效率工具；区别于普通待办清单，"
    "它以日期计划为入口，以专注执行为核心，以贡献热力图形成持续反馈。"
)
set_run_font(r, size=10.3, color=GRAY)

add_heading(doc, "2. 创意功能设计")

status = doc.add_paragraph()
set_spacing(status, after=2)
r = status.add_run("已实现功能（红色为创新特色功能）")
set_run_font(r, size=10.5, bold=True, color=NAVY, font="黑体")

add_bullet(
    doc,
    "日期滑动式任务计划",
    "首页左右滑动即可查看昨天、今天、明天及更远日期；任务可一键移至今日或明日，新建任务自动继承当前浏览日期。",
    innovative=True,
)
add_bullet(
    doc,
    "任务闹钟决策流程",
    "任务可设置可选闹钟；到点响铃并询问是否进入任务。“是”直接进入专注，“等会”5分钟后重响，“关闭”结束提醒。",
    innovative=True,
)
add_bullet(
    doc,
    "Codeforces 风格贡献日历",
    "年度网格按每天完成任务数与专注时长综合色深浅；点击任意日期可查看当天任务、完成状态及专注记录。",
    innovative=True,
)
add_bullet(
    doc,
    "任务与番茄钟闭环",
    "任务仅名称必填，其余采用默认值；支持颜色标签、进度、编辑、完成和删除，并与暂停/继续/提前结束的番茄钟关联。",
)
add_bullet(
    doc,
    "个人主页与可视化统计",
    "本地账户展示学习等级、连续活跃天数、年度专注、完成率、最近7天折线趋势和科目占比。",
)
add_bullet(
    doc,
    "离线数据与隐私保护",
    "任务、专注记录和账户信息保存在本机 SQLite/本地配置中，无需联网；闹钟在系统重启后自动恢复。",
)

future = doc.add_paragraph()
set_spacing(future, before=3, after=2)
r = future.add_run("未来扩展")
set_run_font(r, size=10.5, bold=True, color=NAVY, font="黑体")

add_bullet(doc, "AI 学习助手", "根据历史完成率、任务难度和高效时段自动拆分任务，给出个性化专注时长与学习建议。")
add_bullet(doc, "多端同步与协作", "增加云端备份、跨设备同步、学习小组目标和同伴监督机制。")
add_bullet(doc, "桌面与穿戴设备联动", "提供桌面小组件、快捷开始、智能手表提醒和更丰富的专注环境支持。")

doc.core_properties.title = "Focus Flow 产品设计书"
doc.core_properties.subject = "Android 学习专注与任务管理产品设计"
doc.core_properties.keywords = "Focus Flow, 番茄钟, 任务管理, 贡献热力图, 产品设计"
doc.core_properties.author = ""
doc.core_properties.last_modified_by = ""

doc.save(OUTPUT)
print(OUTPUT)
